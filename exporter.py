import io
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int, color: tuple = None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def _add_post_box(doc: Document, caption: str, label: str = "", extra: str = ""):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "EBF5FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if label:
        run = p.add_run(f"[{label}]  ")
        run.bold = True
        run.font.color.rgb = RGBColor(31, 97, 141)
        run.font.size = Pt(9)
    run = p.add_run(caption)
    run.font.size = Pt(10)
    if extra:
        p2 = cell.add_paragraph()
        run2 = p2.add_run(extra)
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()


def _add_bullet_list(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build_docx(store_data: dict, strategy: dict) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    store_name = store_data.get("store_name", "Online Store")
    today = date.today().strftime("%B %d, %Y")

    # ── Cover ────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(store_name)
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 97, 141)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Social Media Marketing Strategy")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Generated: {today}").font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── 1. Brand Overview ────────────────────────────────────
    _add_heading(doc, "1. Brand Overview", 1, (31, 97, 141))

    brand = strategy.get("brand_overview", {})
    doc.add_paragraph(f"Target Audience: {brand.get('target_audience', '')}")
    doc.add_paragraph(f"Brand Voice: {brand.get('brand_voice', '')}")

    if brand.get("unique_selling_points"):
        _add_heading(doc, "Unique Selling Points", 2)
        _add_bullet_list(doc, brand["unique_selling_points"])

    if brand.get("best_selling_categories"):
        _add_heading(doc, "Key Categories", 2)
        _add_bullet_list(doc, brand["best_selling_categories"])

    # ── 2. Facebook Strategy ─────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "2. Facebook Strategy", 1, (59, 89, 152))

    fb = strategy.get("facebook_strategy", {})

    if fb.get("page_tips"):
        _add_heading(doc, "Page Management Tips", 2)
        _add_bullet_list(doc, fb["page_tips"])

    _add_heading(doc, "Ready-to-Post Captions (copy & paste)", 2)
    for i, post in enumerate(fb.get("posts", []), 1):
        extra_parts = []
        if post.get("post_type"):
            extra_parts.append(f"Type: {post['post_type']}")
        if post.get("best_time"):
            extra_parts.append(f"Best time: {post['best_time']}")
        _add_post_box(
            doc,
            post.get("caption", ""),
            label=f"Post {i}",
            extra="  |  ".join(extra_parts),
        )

    if fb.get("ad_copy_variants"):
        _add_heading(doc, "Facebook Ad Copy Variants", 2)
        for i, ad in enumerate(fb["ad_copy_variants"], 1):
            _add_heading(doc, f"Ad Variant {i}: {ad.get('headline', '')}", 3)
            doc.add_paragraph(ad.get("body", ""))
            p = doc.add_paragraph()
            run = p.add_run(f"Call to Action: {ad.get('call_to_action', '')}")
            run.bold = True

    # ── 3. Instagram Strategy ────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "3. Instagram Strategy", 1, (193, 53, 132))

    ig = strategy.get("instagram_strategy", {})

    _add_heading(doc, "Ready-to-Post Captions (copy & paste)", 2)
    for i, post in enumerate(ig.get("posts", []), 1):
        extra_parts = []
        if post.get("post_type"):
            extra_parts.append(f"Type: {post['post_type']}")
        if post.get("visual_suggestion"):
            extra_parts.append(f"Visual: {post['visual_suggestion']}")
        _add_post_box(
            doc,
            post.get("caption", ""),
            label=f"Post {i}",
            extra="  |  ".join(extra_parts),
        )

    hashtags = ig.get("hashtag_groups", {})
    if hashtags:
        _add_heading(doc, "Hashtag Strategy", 2)
        for group_name, tags in hashtags.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{group_name.capitalize()} reach: ")
            run.bold = True
            p.add_run(" ".join(tags))

    if ig.get("reel_ideas"):
        _add_heading(doc, "Reel Ideas", 2)
        _add_bullet_list(doc, ig["reel_ideas"])

    if ig.get("story_ideas"):
        _add_heading(doc, "Story Ideas", 2)
        _add_bullet_list(doc, ig["story_ideas"])

    # ── 4. 7-Day Content Calendar ─────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "4. 7-Day Content Calendar", 1, (31, 97, 141))

    calendar = strategy.get("content_calendar", [])
    if calendar:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Day", "Theme", "Facebook", "Instagram"]
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            _set_cell_bg(cell, "2C3E50")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        for row_data in calendar:
            row = table.add_row()
            row.cells[0].text = row_data.get("day", "")
            row.cells[1].text = row_data.get("theme", "")
            row.cells[2].text = row_data.get("facebook", "")
            row.cells[3].text = row_data.get("instagram", "")

    # ── 5. Campaigns ─────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "5. Promotional Campaigns", 1, (31, 97, 141))

    for camp in strategy.get("campaigns", []):
        _add_heading(doc, camp.get("name", "Campaign"), 2)
        doc.add_paragraph(camp.get("description", ""))

        info = doc.add_paragraph()
        info.add_run("Duration: ").bold = True
        info.add_run(camp.get("duration", ""))

        fb_p = doc.add_paragraph()
        fb_p.add_run("Facebook approach: ").bold = True
        fb_p.add_run(camp.get("facebook_angle", ""))

        ig_p = doc.add_paragraph()
        ig_p.add_run("Instagram approach: ").bold = True
        ig_p.add_run(camp.get("instagram_angle", ""))

        out_p = doc.add_paragraph()
        out_p.add_run("Expected outcome: ").bold = True
        out_p.add_run(camp.get("expected_outcome", ""))

        doc.add_paragraph()

    # ── 6. Posting Tips ───────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "6. Best Practices & Posting Tips", 1, (31, 97, 141))

    tips = strategy.get("posting_tips", {})
    if tips.get("facebook"):
        _add_heading(doc, "Facebook Tips", 2)
        _add_bullet_list(doc, tips["facebook"])
    if tips.get("instagram"):
        _add_heading(doc, "Instagram Tips", 2)
        _add_bullet_list(doc, tips["instagram"])
    if tips.get("general"):
        _add_heading(doc, "General Tips", 2)
        _add_bullet_list(doc, tips["general"])

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
